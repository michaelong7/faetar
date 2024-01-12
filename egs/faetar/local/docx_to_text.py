#! /usr/bin/env python

import sys
import argparse

from pathlib import Path
from unicodedata import lookup

import docx
import docx.document
from docx.enum.style import WD_STYLE_TYPE


def parse_fmap(pth: Path) -> dict[int, str]:
    fmap = dict()
    with pth.open() as f:
        for no, line in enumerate(f):
            line = line.split(";", maxsplit=1)[0].strip()
            if not line:
                continue
            hex_, name = line.split(maxsplit=1)
            try:
                key = int(hex_, 0)
                val = lookup(name)
            except Exception as e:
                raise IOError(f"Error parsing '{pth}' line {no + 1}:") from e
            fmap[key] = val
    return fmap


def handle_text(text: str, font: str, fmaps: dict[str, dict[int, str]]) -> str:
    fmap = fmaps.get(font, None)
    # sometimes characters < 0x0100 will be stored with a 0xf000 mask. Undo this
    offset = int("0xf000", 0)
    for i in range(int("0xff", 0)):
        text = text.replace(chr(i + offset), chr(i))

    # apply a charater mapping if we've got one
    if fmap is not None:
        text = "".join(map(lambda x: fmap.get(ord(x), x), text))

    # dunno
    if font == "Doulos SIL":
        text = text.replace("ɤ", "j")

    return text


def main(args=None):
    parser = argparse.ArgumentParser(
        description="Parse docx file, outputting text to cmd"
    )
    parser.add_argument("in_file", type=argparse.FileType("rb"))
    parser.add_argument(
        "out_file", nargs="?", type=argparse.FileType("wt"), default=sys.stdout
    )
    parser.add_argument("--map-folder", type=Path, default=None)
    parser.add_argument("--route-through", type=Path, default=None)
    parser.add_argument("--force-font-as", default=None)

    options = parser.parse_args(args)

    route_through = None
    if options.route_through is not None:
        route_through = parse_fmap(options.route_through)

    forced_font = options.force_font_as

    fmaps = dict()
    if options.map_folder is not None:
        map_folder: Path = options.map_folder
        if not map_folder.is_dir():
            raise ValueError(f"{map_folder} is not a directory")
        for pth in map_folder.glob("*.fmap"):
            fmap = parse_fmap(pth)
            if route_through is not None:
                fmap_ = dict()
                for key, val in route_through.items():
                    new_val = fmap.get(key, val)
                    new_key = ord(val)
                    fmap_[new_key] = new_val
                fmap = fmap_
            fmaps[pth.stem] = fmap

    doc: docx.document.Document = docx.Document(options.in_file)

    deft_par_font = doc.styles.default(WD_STYLE_TYPE.PARAGRAPH)
    deft_par_font = None if deft_par_font is None else deft_par_font.font.name
    for par in doc.paragraphs:
        line = ""
        par_font = par.style.font.name
        if par_font is None:
            par_font = deft_par_font
        for run in par.runs:
            font, text = run.font.name, run.text
            if forced_font is not None:
                font = forced_font
            elif font is None:
                font = par_font
            line += handle_text(text, font, fmaps)
        print(line, file=options.out_file)

    deft_tab_font = doc.styles.default(WD_STYLE_TYPE.TABLE)
    if deft_tab_font is None:
        deft_tab_font = deft_par_font  # FIXME(sdrobert): is this right?
    else:
        deft_tab_font = deft_tab_font.font.name
    for tab in doc.tables:
        font = tab.style.font.name
        if forced_font is not None:
            font = forced_font
        elif font is None:
            font = deft_tab_font
        for row in tab.rows:
            line = ""
            for cell in row.cells:
                line += "\t" + handle_text(cell.text, font, fmaps)
            print(line, file=options.out_file)


if __name__ == "__main__":
    sys.exit(main())
